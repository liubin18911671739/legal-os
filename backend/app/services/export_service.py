"""
Report Export Service

This module handles report generation for PDF and DOCX formats.
"""

from typing import Optional, Dict, Any, List
from datetime import datetime
import logging

from PyMuPDF import pdfmetrics
from PyMuPDF.pdf import (
    PageObject,
    PDFWriter,
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.lib import colors
from reportlab.platypus import Paragraph as PLParagraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER, TA_LEFT

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns.qn import CT_Pn

logger = logging.getLogger(__name__)


class ReportGenerator:
    """Generate formatted report content from analysis results"""
    
    def __init__(self):
        """Initialize report generator"""
        pass
    
    def generate_markdown(
        self,
        analysis_result: Dict[str, Any]
    ) -> str:
        """Generate markdown report from analysis result
        
        Args:
            analysis_result: Analysis result from contract analysis
            
        Returns:
            Markdown formatted report
        """
        lines = []
        
        # Header
        lines.append("# 合同审查报告")
        lines.append("")
        lines.append(f"**任务 ID**: {analysis_result.get('task_id', 'N/A')}")
        lines.append(f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Executive Summary
        report = analysis_result.get('report', {})
        lines.append("## 执行摘要")
        lines.append("")
        lines.append(report.get('executive_summary', '无执行摘要'))
        lines.append("")
        lines.append("")
        
        # Overall Risk Assessment
        lines.append("## 整体风险评估")
        overall_risk = analysis_result.get('overall_risk', 'unknown')
        confidence = analysis_result.get('validation_confidence', 0.0)
        
        lines.append(f"**风险等级**: {self._get_risk_badge(overall_risk)}")
        lines.append(f"**验证置信度**: {confidence:.2%}")
        lines.append("")
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Risk Matrix
        lines.append("## 风险矩阵")
        lines.append("")
        risk_matrix = report.get('risk_matrix', {})
        lines.append("| 风险类型 | 风险等级 |")
        lines.append("| --- | --- |")
        
        risk_levels = {'low': '低', 'medium': '中', 'high': '高'}
        risk_fields = {
            'legal_risk': '法律风险',
            'financial_risk': '财务风险',
            'operational_risk': '运营风险',
            'strategic_risk': '战略风险'
        }
        
        for field_key, label in risk_fields.items():
            risk_value = risk_matrix.get(field_key, 'N/A')
            risk_level = risk_levels.get(risk_value, 'N/A')
            lines.append(f"| {label} | {risk_level} |")
        
        lines.append("")
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Findings
        lines.append("## 发现的问题")
        lines.append("")
        
        findings = report.get('findings', [])
        if not findings:
            lines.append("未发现问题")
        else:
            for i, finding in enumerate(findings, 1):
                severity = finding.get('severity', 'unknown')
                category = finding.get('category', 'general')
                lines.append(f"### {i}. {self._get_severity_badge(severity)} {category}")
                lines.append("")
                lines.append(f"**描述**: {finding.get('description', '无描述')}")
                lines.append("")
                
                if 'clause_reference' in finding:
                    lines.append(f"**条款引用**: 第 {finding['clause_reference']} 条")
                    lines.append("")
                
                if 'suggestion' in finding:
                    lines.append(f"**建议**: {finding['suggestion']}")
                    lines.append("")
                
                if 'citation' in finding:
                    lines.append(f"**引用**: {finding['citation']}")
                    lines.append("")
                
                lines.append("---")
                lines.append("")
        
        # Suggestions
        lines.append("## 建议")
        lines.append("")
        
        suggestions = report.get('suggestions', [])
        if not suggestions:
            lines.append("无具体建议")
        else:
            for i, suggestion in enumerate(suggestions, 1):
                lines.append(f"{i}. {suggestion}")
                lines.append("")
        
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Agent History
        lines.append("## 智能体执行历史")
        lines.append("")
        
        agent_history = analysis_result.get('agent_history', [])
        if not agent_history:
            lines.append("无智能体执行历史")
        else:
            for agent in agent_history:
                lines.append(f"- {agent}")
            lines.append("")
        
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Footer
        lines.append("---")
        lines.append("")
        lines.append(f"*本报告由 LegalOS AI 系统自动生成于 {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")
        
        return "\n".join(lines)
    
    def _get_risk_badge(self, risk: str) -> str:
        """Get risk badge markdown"""
        risk_badges = {
            'low': '<span style="color:green">低风险</span>',
            'medium': '<span style="color:orange">中风险</span>',
            'high': '<span style="color:red">高风险</span>',
            'unknown': '<span style="color:gray">未知</span>'
        }
        return risk_badges.get(risk, risk_badges['unknown'])
    
    def _get_severity_badge(self, severity: str) -> str:
        """Get severity badge markdown"""
        severity_badges = {
            'low': '<span style="color:green">低</span>',
            'medium': '<span style="color:orange">中</span>',
            'high': '<span style="color:red">高</span>',
            'critical': '<span style="color:darkred">严重</span>',
            'unknown': '<span style="gray">未知</span>'
        }
        return severity_badges.get(severity, severity_badges['unknown'])


class PDFExportService:
    """Export report to PDF format using PyMuPDF"""
    
    def __init__(self):
        """Initialize PDF export service"""
        pass
    
    def export_to_pdf(
        self,
        markdown_content: str,
        title: str = "合同审查报告"
    ) -> bytes:
        """Export markdown content to PDF
        
        Args:
            markdown_content: Markdown formatted report
            title: Report title
            
        Returns:
            PDF file as bytes
        """
        try:
            logger.info(f"Generating PDF: {title}")
            
            # Create PDF
            buffer = self._create_pdf_from_markdown(
                markdown_content,
                title
            )
            
            logger.info(f"PDF generated successfully, size: {len(buffer)} bytes")
            return buffer
            
        except Exception as e:
            logger.error(f"PDF generation failed: {e}", exc_info=True)
            raise
    
    def _create_pdf_from_markdown(
        self,
        markdown_content: str,
        title: str
    ) -> bytes:
        """Create PDF from markdown content"""
        # Create PDF writer
        buffer = BytesIO()
        doc = SimpleDocTemplate("Arial", 10, title)
        doc.pagesize = letter
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        heading1_style = styles['Heading1']
        heading2_style = styles['Heading2']
        normal_style = styles['BodyText']
        code_style = styles['Code']
        
        # Parse markdown and create PDF elements
        lines = markdown_content.split('\n')
        
        y_position = 700
        for line in lines:
            line = line.strip()
            
            if not line:
                doc.append(Spacer(1, 0.2*12))
                continue
            
            # Headers
            if line.startswith('# '##'):
                doc.add_page_break()
                doc.add(Paragraph(line.replace('##', ''), heading2_style))
                y_position = 750
            elif line.startswith('# ''):
                doc.add_page_break()
                doc.add(Paragraph(line.replace('#', ''), heading1_style)
                y_position = 750
            # Bold text
            elif line.startswith('**') and line.endswith('**'):
                text = line.replace('**', '')
                doc.add(Paragraph(text, normal_style))
                y_position += 14
            # Code/inline code
            elif line.startswith('|'):
                # Parse as table
                doc.add(self._parse_table_line(line))
                y_position += 14
            # Normal text
            elif line.startswith('- ''):
                text = line.lstrip('- ').strip()
                doc.add(Paragraph(f"• {text}", normal_style))
                y_position += 12
            # Horizontal rule
            elif line.startswith('---'):
                doc.add(Spacer(1, 0.2*12))
                y_position += 12
            # List items
            elif line.startswith(f"{i}."):
                text = line.replace(f"{i}. ", "").strip()
                doc.add(Paragraph(text, normal_style))
                y_position += 12
            # HTML/markdown badges
            elif '<span style=' in line:
                doc.add_paragraph(htmlparser.HTMLParser().parse(line))
                y_position += 12
            else:
                doc.add(Paragraph(line, normal_style))
                y_position += 12
            
            # Add page break if needed
            if y_position > 700:
                doc.add_page_break()
                y_position = 100
        
        # Footer
        doc.add_page_break()
        doc.add Paragraph(f"Generated by LegalOS at {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style)
        
        # Build PDF
        doc.build(buffer)
        
        return buffer.getvalue()


class DOCXExportService:
    """Export report to DOCX format using python-docx"""
    
    def __init__(class DOCXExportService):
        """Initialize DOCX export service"""
        pass
    
    def export_to_docx(
        self,
        markdown_content: str,
        title: str = "合同审查报告"
    ) -> bytes:
        """Export markdown content to DOCX
        
        Args:
            markdown_content: Markdown formatted report
            title: Report title
            
        Returns:
            DOCX file as bytes
        """
        try:
            logger.info(f"Generating DOCX: {title}")
            
            # Create Word document
            doc = Document()
            
            # Add title
            title_paragraph = doc.add_heading(
                title,
                level=0
            )
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH
            
            # Set document style
            style = doc.styles['Normal']
            font = doc.styles['Normal Font']
            font.size = Pt(11)
            
            # Parse markdown and add to DOCX
            lines = markdown_content.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    doc.add_paragraph('')
                    continue
                
                # Headers
                if line.startswith('## '):
                    doc.add_heading(line.replace('##', ''), level=1)
                elif line.startswith('# '):
                    doc.add_heading(line.replace('#', ''), level=0)
                # Bold text
                elif line.startswith('**') and line.endswith('**'):
                    text = line.replace('**', '')
                    paragraph = doc.add_paragraph(text)
                    run = paragraph.add_run(bold=True)
                # List items
                elif line.startswith('- '):
                    text = line.lstrip('- ').strip()
                    doc.add_paragraph(text, style=style)
                # Horizontal rule
                elif line.startswith('---'):
                    doc.add_paragraph('_' * 50)
                # Numbered list
                elif line.startswith(f"{i}."):
                    text = line.replace(f"{i}. ", "").strip()
                    doc.add_paragraph(text, style=style)
                else:
                    # Regular paragraph
                    doc.add_paragraph(line, style=style)
            
            # Add footer
            doc.add_page_break()
            footer = doc.paragraph(f"Generated by LegalOS at {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style=style)
            footer.alignment = WD_ALIGN_PARAGRAPH
            
            # Save to bytes
            buffer = BytesIO()
            doc.save(buffer)
            
            logger.info(f"DOCX generated successfully, size: {len(buffer.getvalue())} bytes")
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}", exc_info=True)
            raise


class ExportService:
    """Main export service coordinating PDF and DOCX exports"""
    
    def __init__(self):
        """Initialize export service"""
        self.report_generator = ReportGenerator()
        self.pdf_exporter = PDFExportService()
        self.docx_exporter = DOCXExportService()
    
    def export_report(
        self,
        analysis_result: Dict[str, Any],
        format: str = 'pdf',
        title: str = "合同审查报告"
    ) -> bytes:
        """Export analysis report to specified format
        
        Args:
            analysis_result: Analysis result from contract analysis
            format: Export format ('pdf' or 'docx')
            title: Report title
            
        Returns:
            File content as bytes
        """
        # Validate format
        if format not in ['pdf', 'docx']:
            raise ValueError(f"Unsupported format: {format}")
        
        # Generate markdown
        markdown_content = self.report_generator.generate_markdown(analysis_result)
        
        # Export to requested format
        if format == 'pdf':
            return self.pdf_exporter.export_to_pdf(markdown_content, title)
        elif format == 'docx':
            return self.docx_exporter.export_to_docx(markdown_content, title)
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    def export_report_with_charts(
        self,
        analysis_result: Dict[str, Any],
        evaluation_results: Optional[List[Dict[str, Any]]] = None,
        format: str = 'pdf',
        title: str = "合同审查报告"
    ) -> bytes:
        """Export report with embedded charts
        
        Args:
            analysis_result: Single contract analysis result
            evaluation_results: Comparison results from multiple baselines
            format: Export format
            title: Report title
            
        Returns:
            File content as bytes
        """
        # Generate markdown with charts section
        markdown_content = self.report_generator.generate_markdown(analysis_result)
        
        if evaluation_results:
            markdown_content += "\n\n"
            markdown_content += "## 评估结果对比\n\n"
            markdown_content += self._generate_evaluation_comparison_markdown(evaluation_results)
        
        # Export to requested format
        return self.export_report(analysis_result, format, title)
    
    def _generate_evaluation_comparison_markdown(
        self,
        evaluation_results: List[Dict[str, Any]]
    ) -> str:
        """Generate markdown for evaluation comparison"""
        lines = []
        
        lines.append("### F1 分数对比")
        lines.append("")
        lines.append("| 基线类型 | F1 分数 | 幻觉率 | 耗时 | 成本 |")
        lines.append("| --- | --- | --- | --- | --- |")
        
        for result in evaluation_results:
            baseline = result.get('baseline_type', 'unknown')
            metrics = result.get('metrics', {})
            f1 = metrics.get('f1_score', 0)
            hallucination = metrics.get('hallucination_rate', 0)
            duration = result.get('duration', 0)
            cost = result.get('cost', 0)
            
            lines.append(f"| {baseline} | {f1:.2%} | {hallucination:.2%} | {duration:.1f} | ${cost:.2f} |")
        
        lines.append("")
        
        return "\n".join(lines)


# Helper imports
from io import BytesIO
import html.parser as htmlparser
from reportlab.lib.pagesizes import letter


def create_export_service() -> ExportService:
    """Factory function to create export service"""
    return ExportService()
