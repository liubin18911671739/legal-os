from docx import Document
from typing import Optional, List
from app.rag.loaders.base_loader import BaseLoader, Document as RagDocument, FileType


class DocxLoader(BaseLoader):
    """DOCX document loader using python-docx."""
    
    def load(self, file_path: str) -> RagDocument:
        """Load DOCX document and extract text."""
        self.validate_file(file_path)
        
        # Extract metadata
        metadata = self.extract_metadata(file_path)
        
        # Extract text content
        text_content = self._extract_text(file_path)
        
        # Create document object
        return RagDocument(
            id=metadata["id"],
            title=metadata["title"],
            file_name=metadata["file_name"],
            file_type=FileType.DOCX,
            content=text_content,
            file_path=file_path,
            file_size=metadata["file_size"],
            metadata=metadata,
        )
    
    def _extract_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            
            # Extract paragraphs from all runs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Add section markers
            full_text = "\\n\\n".join(paragraphs)
            
            return full_text
        except Exception as e:
            print(f"Error loading DOCX: {e}")
            return f"Error loading document: {str(e)}"
